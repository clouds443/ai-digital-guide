# -*- coding: utf-8 -*-
import importlib.util
import json
import os
import threading
import time
import uuid
from io import BytesIO

from runtime_paths import asset_path, runtime_path

ROOT_DIR = asset_path()
CACHE_DIR = runtime_path(".cache")
EVALUATOR_PATH = asset_path("scripts", "evaluate_lingshan_qa.py")
HIGH_RISK_CASE_IDS = {"Q021", "Q022", "Q023", "Q027", "Q051", "Q058", "Q086"}
DEEPSEEK_MIXED_CACHE_MESSAGE = "旧缓存包含规则直答，请重新运行 DeepSeek 评测。"
EVALUATION_JOBS = {}
EVALUATION_JOBS_LOCK = threading.Lock()
ACTIVE_JOB_STATUSES = {"queued", "running", "reviewing", "cancelling"}


def _normalize_mode(mode):
    value = str(mode or "deepseek").strip().lower()
    return value if value in {"deepseek", "local"} else "deepseek"


def _cache_path(mode):
    return os.path.join(CACHE_DIR, "lingshan_eval_{0}.json".format(_normalize_mode(mode)))


def _latest_cache_path(mode):
    return os.path.join(CACHE_DIR, "lingshan_eval_{0}_latest.json".format(_normalize_mode(mode)))


def _load_evaluator():
    spec = importlib.util.spec_from_file_location("evaluate_lingshan_qa", EVALUATOR_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _read_json(path, default):
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def _write_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp_path = path + ".tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, path)


def _as_text_list(value):
    if not value:
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    return [str(value)]


def _score_value(item):
    try:
        return float(item.get("score", 0))
    except Exception:
        return 0.0


def _current_scoring_version():
    try:
        evaluator = _load_evaluator()
        return getattr(evaluator, "SCORING_VERSION", "relaxed_keyword_v3")
    except Exception:
        return "relaxed_keyword_v3"


def _normalize_case_item(item):
    raw = dict(item or {})
    answer = str(raw.get("answer", ""))
    expected = dict(raw.get("expected") or {})
    must_include = _as_text_list(expected.get("must_include") or raw.get("must_include"))
    must_not_include = _as_text_list(expected.get("must_not_include") or raw.get("must_not_include"))
    matched_include = _as_text_list((raw.get("hit_detail") or {}).get("matched_include"))
    missing_include = _as_text_list((raw.get("hit_detail") or {}).get("missing_include"))
    forbidden_hits = _as_text_list((raw.get("hit_detail") or {}).get("forbidden_hits") or raw.get("forbidden_hits"))
    if must_include and not matched_include and not missing_include:
        matched_include = [word for word in must_include if word in answer]
        missing_include = [word for word in must_include if word not in answer]
    if must_not_include and not forbidden_hits:
        forbidden_hits = [word for word in must_not_include if word in answer]
    hit_detail = {
        "include_hits": len(matched_include),
        "include_total": len(must_include),
        "matched_include": matched_include,
        "missing_include": missing_include,
        "forbidden_hits": forbidden_hits,
    }
    normalized = {
        "id": str(raw.get("id", "")),
        "category": str(raw.get("category", "")),
        "question": str(raw.get("question", "")),
        "expected": {
            "must_include": must_include,
            "must_not_include": must_not_include,
            "source_doc": str(expected.get("source_doc") or raw.get("source_doc") or ""),
            "weight": expected.get("weight", raw.get("weight", 1)),
        },
        "answer": answer,
        "answer_provider": str(raw.get("answer_provider", "")),
        "score": _score_value(raw),
        "hit_detail": hit_detail,
        "failure_types": _as_text_list(raw.get("failure_types")),
        "latency_ms": int(raw.get("latency_ms") or 0),
        "sources": raw.get("sources") or [],
    }
    if isinstance(raw.get("semantic_review"), dict):
        normalized["semantic_review"] = dict(raw.get("semantic_review") or {})
    return normalized


def _build_category_stats(items):
    grouped = {}
    for item in items:
        category = item.get("category") or "uncategorized"
        bucket = grouped.setdefault(category, {"category": category, "total": 0, "passed": 0, "low_score": 0, "score_sum": 0.0})
        score = _score_value(item)
        bucket["total"] += 1
        bucket["score_sum"] += score
        if score >= 8:
            bucket["passed"] += 1
        else:
            bucket["low_score"] += 1
    stats = []
    for category in sorted(grouped.keys()):
        bucket = grouped[category]
        total = bucket["total"] or 1
        stats.append({
            "category": category,
            "total": bucket["total"],
            "passed": bucket["passed"],
            "low_score": bucket["low_score"],
            "avg_score": round(bucket["score_sum"] / total, 2),
            "pass_rate": round(bucket["passed"] / float(total), 4),
        })
    return stats


def compact_evaluation_summary(summary):
    payload = dict(summary or {})
    items = [_normalize_case_item(item) for item in (payload.get("case_items") or payload.get("items") or [])]
    low_score_items = [item for item in items if float(item.get("score", 0)) < 8]
    high_risk_failures = [
        item for item in items
        if item.get("id") in HIGH_RISK_CASE_IDS and float(item.get("score", 0)) < 8
    ]
    payload["case_items"] = items
    payload["category_stats"] = _build_category_stats(items)
    payload["provider_stats"] = payload.get("provider_stats") or _build_provider_stats(items)
    payload["low_score_items"] = low_score_items[:12]
    payload["high_risk_failures"] = high_risk_failures
    payload["ready"] = True
    payload["ok"] = True
    payload["updated_at"] = payload.get("updated_at") or time.strftime("%Y-%m-%d %H:%M:%S")
    payload.pop("items", None)
    return payload


def _build_provider_stats(items):
    stats = {}
    for item in items:
        provider = item.get("answer_provider") or "unknown"
        stats[provider] = stats.get(provider, 0) + 1
    return stats


def _recalculate_payload_summary(payload, items):
    next_payload = dict(payload or {})
    normalized_items = [_normalize_case_item(item) for item in (items or [])]
    total = 0.0
    max_score = 0.0
    fact_hits = 0
    for item in normalized_items:
        expected = item.get("expected") or {}
        try:
            weight = float(expected.get("weight", 1) or 1)
        except Exception:
            weight = 1.0
        hit = item.get("hit_detail") or {}
        total += _score_value(item) * weight
        max_score += 10.0 * weight
        if hit.get("include_hits") == hit.get("include_total") and not hit.get("forbidden_hits"):
            fact_hits += 1
    next_payload["case_items"] = normalized_items
    next_payload.pop("items", None)
    next_payload["case_count"] = len(normalized_items)
    next_payload["score_percent"] = round(total / max_score * 100, 2) if max_score else 0.0
    next_payload["fact_accuracy"] = round(fact_hits / float(len(normalized_items) or 1), 4)
    next_payload["failed_count"] = len([item for item in normalized_items if _score_value(item) < 8])
    next_payload["avg_latency_ms"] = round(
        sum(int(item.get("latency_ms") or 0) for item in normalized_items) / float(len(normalized_items) or 1),
        2,
    )
    next_payload["provider_stats"] = _build_provider_stats(normalized_items)
    next_payload["low_score_items"] = [item for item in normalized_items if _score_value(item) < 8][:12]
    next_payload["high_risk_failures"] = [
        item for item in normalized_items
        if item.get("id") in HIGH_RISK_CASE_IDS and _score_value(item) < 8
    ]
    next_payload["category_stats"] = _build_category_stats(normalized_items)
    next_payload["ready"] = True
    next_payload["ok"] = True
    next_payload["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    return next_payload


def _rescore_evaluation_payload(payload):
    if not payload or not payload.get("ready"):
        return payload, False
    version = _current_scoring_version()
    if payload.get("scoring_version") == version:
        return payload, False
    raw_items = payload.get("case_items") or payload.get("items") or []
    if not raw_items:
        return payload, False
    evaluator = _load_evaluator()
    items = []
    total = 0.0
    max_score = 0.0
    fact_hits = 0
    for raw_item in raw_items:
        item = _normalize_case_item(raw_item)
        expected = item.get("expected") or {}
        case = {
            "id": item.get("id", ""),
            "category": item.get("category", ""),
            "question": item.get("question", ""),
            "must_include": _as_text_list(expected.get("must_include")),
            "must_not_include": _as_text_list(expected.get("must_not_include")),
            "source_doc": expected.get("source_doc", ""),
            "weight": expected.get("weight", 1),
        }
        scored = evaluator.score_case(case, item.get("answer", ""))
        item["score"] = scored["score"]
        item["hit_detail"] = {
            "include_hits": scored["include_hits"],
            "include_total": scored["include_total"],
            "matched_include": scored["matched_include"],
            "missing_include": scored["missing_include"],
            "forbidden_hits": scored["forbidden_hits"],
        }
        item["failure_types"] = scored["failure_types"]
        items.append(item)
        weight = float(expected.get("weight", 1) or 1)
        total += float(scored["score"]) * weight
        max_score += 10.0 * weight
        if scored["include_hits"] == scored["include_total"] and not scored["forbidden_hits"]:
            fact_hits += 1
    rescored = dict(payload)
    rescored["case_items"] = items
    rescored.pop("items", None)
    rescored["case_count"] = int(rescored.get("case_count") or len(items))
    rescored["score_percent"] = round(total / max_score * 100, 2) if max_score else 0.0
    rescored["fact_accuracy"] = round(fact_hits / float(len(items) or 1), 4)
    rescored["failed_count"] = len([item for item in items if float(item.get("score", 0)) < 8])
    rescored["avg_latency_ms"] = round(sum(int(item.get("latency_ms") or 0) for item in items) / float(len(items) or 1), 2)
    rescored["provider_stats"] = rescored.get("provider_stats") or _build_provider_stats(items)
    rescored["low_score_items"] = [item for item in items if float(item.get("score", 0)) < 8][:12]
    rescored["high_risk_failures"] = [
        item for item in items
        if item.get("id") in HIGH_RISK_CASE_IDS and float(item.get("score", 0)) < 8
    ]
    rescored["category_stats"] = _build_category_stats(items)
    rescored["scoring_version"] = version
    return rescored, True


def _deepseek_cache_invalid_message(payload):
    if not payload or payload.get("mode") != "deepseek" or not payload.get("ready"):
        return ""
    items = payload.get("case_items") or payload.get("items") or []
    stats = dict(payload.get("provider_stats") or {})
    if not stats and items:
        stats = _build_provider_stats(items)
    if not stats:
        return ""
    try:
        case_count = int(payload.get("case_count") or len(items) or sum(int(v or 0) for v in stats.values()))
    except Exception:
        case_count = len(items) or 0
    deepseek_count = int(stats.get("deepseek") or 0)
    non_deepseek = [key for key, value in stats.items() if key != "deepseek" and value]
    if non_deepseek or (case_count and deepseek_count != case_count):
        return DEEPSEEK_MIXED_CACHE_MESSAGE
    return ""


def _with_deepseek_cache_validation(payload):
    message = _deepseek_cache_invalid_message(payload)
    if not message:
        return payload
    invalid = dict(payload or {})
    invalid.update({
        "ok": False,
        "ready": False,
        "invalid_cache": True,
        "message": message,
        "error": message,
    })
    return invalid


def _recover_legacy_snapshot(mode, cached):
    latest = _read_json(_latest_cache_path(mode), None)
    if latest and (latest.get("items") or latest.get("case_items")):
        payload = compact_evaluation_summary(latest)
        cached = dict(cached or {})
        if cached:
            for key in ["updated_at"]:
                if cached.get(key):
                    payload[key] = cached[key]
        _write_json(_cache_path(mode), payload)
        return payload
    return cached


def get_evaluation_snapshot(mode="deepseek"):
    mode = _normalize_mode(mode)
    cached = _read_json(_cache_path(mode), None)
    if cached:
        if cached.get("ready") and not cached.get("case_items"):
            cached = _recover_legacy_snapshot(mode, cached)
        cached, rescored = _rescore_evaluation_payload(cached)
        if rescored:
            _write_json(_cache_path(mode), cached)
        return _with_deepseek_cache_validation(cached)
    return {
        "ok": False,
        "ready": False,
        "mode": mode,
        "model": "deepseek-chat" if mode == "deepseek" else "local-rag",
        "message": "尚未运行 {0} 评测。".format("DeepSeek" if mode == "deepseek" else "本地基线"),
    }


def latest_evaluation_summary():
    deepseek = get_evaluation_snapshot("deepseek")
    if deepseek.get("ready"):
        return deepseek
    local = get_evaluation_snapshot("local")
    return local if local.get("ready") else deepseek


def _execute_evaluation(mode="deepseek", progress_callback=None, cancel_checker=None):
    mode = _normalize_mode(mode)
    evaluator = _load_evaluator()
    summary = evaluator.evaluate(mode=mode, progress_callback=progress_callback, cancel_checker=cancel_checker)
    full_payload = dict(summary or {})
    full_payload["updated_at"] = full_payload.get("updated_at") or time.strftime("%Y-%m-%d %H:%M:%S")
    payload = compact_evaluation_summary(summary)
    invalid_message = _deepseek_cache_invalid_message(payload)
    if invalid_message:
        raise ValueError(invalid_message)
    _write_json(_latest_cache_path(mode), full_payload)
    _write_json(_cache_path(mode), payload)
    return payload


def _execute_semantic_review(mode="deepseek", progress_callback=None, cancel_checker=None, semantic_reviewer=None):
    mode = _normalize_mode(mode)
    if mode != "deepseek":
        raise ValueError("语义复核仅支持 DeepSeek 评测。")
    snapshot = get_evaluation_snapshot(mode)
    if snapshot.get("invalid_cache"):
        raise ValueError(snapshot.get("message") or DEEPSEEK_MIXED_CACHE_MESSAGE)
    if not snapshot.get("ready") or not snapshot.get("case_items"):
        raise ValueError("请先运行 DeepSeek 评测后再进行低分题语义复核。")
    evaluator = _load_evaluator()
    if not hasattr(evaluator, "semantic_review_low_score_items"):
        raise ValueError("当前评测脚本不支持低分题语义复核。")
    items, stats = evaluator.semantic_review_low_score_items(
        snapshot.get("case_items") or [],
        reviewer=semantic_reviewer,
        progress_callback=progress_callback,
        cancel_checker=cancel_checker,
        mode=mode,
    )
    payload = _recalculate_payload_summary(snapshot, items)
    payload["semantic_review_version"] = getattr(evaluator, "SEMANTIC_REVIEW_VERSION", "llm_low_score_review_v1")
    payload["semantic_review_stats"] = stats
    payload["scoring_version"] = payload.get("scoring_version") or _current_scoring_version()
    invalid_message = _deepseek_cache_invalid_message(payload)
    if invalid_message:
        raise ValueError(invalid_message)
    _write_json(_latest_cache_path(mode), payload)
    _write_json(_cache_path(mode), payload)
    return payload


def _progress_payload(job):
    payload = dict(job or {})
    payload.pop("thread", None)
    started_at = float(payload.pop("started_at", 0) or 0)
    if started_at:
        payload["elapsed_ms"] = int((time.time() - started_at) * 1000)
    else:
        payload["elapsed_ms"] = int(payload.get("elapsed_ms") or 0)
    total = int(payload.get("total") or 0)
    completed = int(payload.get("completed") or 0)
    if total:
        payload["percent"] = int(round(completed / float(total) * 100))
    else:
        payload["percent"] = int(payload.get("percent") or 0)
    return payload


def _find_running_job(mode):
    for job in EVALUATION_JOBS.values():
        if job.get("mode") == mode and job.get("status") in ACTIVE_JOB_STATUSES:
            return job
    return None


def _safe_client_job_id(mode, job_id):
    value = str(job_id or "").strip()
    prefix = "eval-{0}-client-".format(mode)
    allowed = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_.")
    if value.startswith(prefix) and len(value) <= 96 and all(ch in allowed for ch in value):
        return value
    return ""


def _is_stale_evaluation_job(job):
    if not job or job.get("status") not in ACTIVE_JOB_STATUSES:
        return False
    thread = job.get("thread")
    if thread is not None:
        try:
            return not thread.is_alive()
        except Exception:
            return True
    started_at = float(job.get("started_at") or 0)
    return bool(started_at and time.time() - started_at > 10)


def _expire_stale_evaluation_jobs(mode):
    for job in EVALUATION_JOBS.values():
        if job.get("mode") != mode:
            continue
        if not _is_stale_evaluation_job(job):
            continue
        job["status"] = "failed"
        job["error"] = "评测任务已停止响应，请重新运行。"
        job["message"] = "评测任务已停止响应，请重新运行。"
        job["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")


def _update_job(job_id, **updates):
    with EVALUATION_JOBS_LOCK:
        job = EVALUATION_JOBS.get(job_id)
        if not job:
            return None
        job.update(updates)
        job["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        return _progress_payload(job)


def _job_cancel_requested(job_id):
    with EVALUATION_JOBS_LOCK:
        job = EVALUATION_JOBS.get(job_id)
        return bool(job and job.get("cancel_requested"))


def _apply_progress_event(job_id, event):
    data = dict(event or {})
    cancelling = _job_cancel_requested(job_id)
    event_name = data.get("event", "")
    status = "reviewing" if str(event_name).startswith("semantic_review") else "running"
    updates = {
        "status": "cancelling" if cancelling else status,
        "total": data.get("total", 0),
        "completed": data.get("completed", 0),
        "current_case_id": data.get("current_case_id", ""),
        "current_question": data.get("current_question", ""),
        "message": "正在结束评测，当前题完成后停止。" if cancelling else data.get("message", ""),
        "error": data.get("error", ""),
    }
    if data.get("event") == "cancelled":
        updates["status"] = "cancelled"
        updates["message"] = data.get("message") or "评测已由管理员终止。"
    if data.get("event") == "retrying":
        updates["message"] = data.get("message") or "正在重试 {0}".format(data.get("current_case_id", ""))
    return _update_job(job_id, **updates)


def _is_cancelled_exception(exc):
    return exc.__class__.__name__ == "EvaluationCancelled"


def _run_evaluation_job(job_id):
    with EVALUATION_JOBS_LOCK:
        job = EVALUATION_JOBS.get(job_id)
        if not job:
            return
        mode = job.get("mode", "deepseek")
        job["status"] = "running"
        job["started_at"] = time.time()
        job["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    try:
        payload = _execute_evaluation(
            mode,
            progress_callback=lambda event: _apply_progress_event(job_id, event),
            cancel_checker=lambda: _job_cancel_requested(job_id),
        )
        _update_job(
            job_id,
            status="completed",
            completed=int(payload.get("case_count") or 0),
            total=int(payload.get("case_count") or 0),
            current_case_id="",
            current_question="",
            message="评测完成",
            error="",
            evaluation=payload,
        )
    except Exception as exc:
        if _is_cancelled_exception(exc):
            _update_job(
                job_id,
                status="cancelled",
                message="评测已由管理员终止，未写入成功缓存。",
                error="",
                cancel_requested=True,
            )
            return
        previous = get_evaluation_snapshot(mode)
        updates = {
            "status": "failed",
            "error": str(exc),
            "message": str(exc),
        }
        if previous and previous.get("ready"):
            updates["evaluation"] = previous
        _update_job(job_id, **updates)


def _run_semantic_review_job(job_id):
    with EVALUATION_JOBS_LOCK:
        job = EVALUATION_JOBS.get(job_id)
        if not job:
            return
        mode = job.get("mode", "deepseek")
        job["status"] = "reviewing"
        job["started_at"] = time.time()
        job["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    try:
        payload = _execute_semantic_review(
            mode,
            progress_callback=lambda event: _apply_progress_event(job_id, event),
            cancel_checker=lambda: _job_cancel_requested(job_id),
        )
        stats = payload.get("semantic_review_stats") or {}
        _update_job(
            job_id,
            status="completed",
            completed=int(stats.get("reviewed") or 0),
            total=int(stats.get("reviewed") or 0),
            current_case_id="",
            current_question="",
            message="低分题语义复核完成",
            error="",
            evaluation=payload,
        )
    except Exception as exc:
        if _is_cancelled_exception(exc):
            _update_job(
                job_id,
                status="cancelled",
                message="低分题语义复核已由管理员终止，未写入成功缓存。",
                error="",
                cancel_requested=True,
            )
            return
        previous = get_evaluation_snapshot(mode)
        updates = {
            "status": "failed",
            "error": str(exc),
            "message": str(exc),
        }
        if previous and previous.get("ready"):
            updates["evaluation"] = previous
        _update_job(job_id, **updates)


def start_evaluation_job(mode="deepseek", job_id=None):
    mode = _normalize_mode(mode)
    with EVALUATION_JOBS_LOCK:
        _expire_stale_evaluation_jobs(mode)
        requested_job_id = _safe_client_job_id(mode, job_id)
        if requested_job_id and requested_job_id in EVALUATION_JOBS:
            progress = _progress_payload(EVALUATION_JOBS[requested_job_id])
            return {"job_id": progress["job_id"], "progress": progress}
        existing = _find_running_job(mode)
        if existing:
            progress = _progress_payload(existing)
            return {"job_id": progress["job_id"], "progress": progress}
        job_id = requested_job_id or "eval-{0}-{1}-{2}".format(mode, int(time.time() * 1000), uuid.uuid4().hex[:8])
        thread = threading.Thread(target=_run_evaluation_job, args=(job_id,), daemon=True)
        job = {
            "job_id": job_id,
            "mode": mode,
            "status": "running",
            "total": 100,
            "completed": 0,
            "percent": 0,
            "current_case_id": "",
            "current_question": "",
            "message": "正在启动评测任务",
            "error": "",
            "cancel_requested": False,
            "started_at": time.time(),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "thread": thread,
        }
        EVALUATION_JOBS[job_id] = job
        progress = _progress_payload(job)
    try:
        thread.start()
    except Exception as exc:
        progress = _update_job(
            job_id,
            status="failed",
            error="评测任务启动失败：{0}".format(exc),
            message="评测任务启动失败：{0}".format(exc),
        )
    return {"job_id": job_id, "progress": progress}


def start_semantic_review_job(mode="deepseek", job_id=None):
    mode = _normalize_mode(mode)
    if mode != "deepseek":
        raise ValueError("语义复核仅支持 DeepSeek 评测。")
    with EVALUATION_JOBS_LOCK:
        _expire_stale_evaluation_jobs(mode)
        requested_job_id = _safe_client_job_id(mode, job_id)
        if requested_job_id and requested_job_id in EVALUATION_JOBS:
            progress = _progress_payload(EVALUATION_JOBS[requested_job_id])
            return {"job_id": progress["job_id"], "progress": progress}
        existing = _find_running_job(mode)
        if existing:
            progress = _progress_payload(existing)
            return {"job_id": progress["job_id"], "progress": progress}
        job_id = requested_job_id or "review-{0}-{1}-{2}".format(mode, int(time.time() * 1000), uuid.uuid4().hex[:8])
        thread = threading.Thread(target=_run_semantic_review_job, args=(job_id,), daemon=True)
        job = {
            "job_id": job_id,
            "mode": mode,
            "kind": "semantic_review",
            "status": "reviewing",
            "total": 0,
            "completed": 0,
            "percent": 0,
            "current_case_id": "",
            "current_question": "正在启动低分题语义复核",
            "message": "正在启动低分题语义复核",
            "error": "",
            "cancel_requested": False,
            "started_at": time.time(),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "thread": thread,
        }
        EVALUATION_JOBS[job_id] = job
        progress = _progress_payload(job)
    try:
        thread.start()
    except Exception as exc:
        progress = _update_job(
            job_id,
            status="failed",
            error="低分题语义复核任务启动失败：{0}".format(exc),
            message="低分题语义复核任务启动失败：{0}".format(exc),
        )
    return {"job_id": job_id, "progress": progress}


def cancel_evaluation_job(mode="deepseek", job_id=None):
    mode = _normalize_mode(mode)
    with EVALUATION_JOBS_LOCK:
        requested_job_id = str(job_id or "")
        job = EVALUATION_JOBS.get(requested_job_id)
        if requested_job_id and not job:
            return {
                "job_id": requested_job_id,
                "mode": mode,
                "status": "not_found",
                "total": 0,
                "completed": 0,
                "percent": 0,
                "current_case_id": "",
                "current_question": "",
                "elapsed_ms": 0,
                "message": "未找到正在运行的评测任务。",
                "error": "",
                "cancel_requested": False,
            }
        if not job:
            job = _find_running_job(mode)
        if not job:
            return {
                "job_id": requested_job_id,
                "mode": mode,
                "status": "not_found",
                "total": 0,
                "completed": 0,
                "percent": 0,
                "current_case_id": "",
                "current_question": "",
                "elapsed_ms": 0,
                "message": "未找到正在运行的评测任务。",
                "error": "",
                "cancel_requested": False,
            }
        if job.get("status") in {"completed", "failed", "cancelled"}:
            return _progress_payload(job)
        job["cancel_requested"] = True
        job["status"] = "cancelling"
        job["message"] = "正在结束评测，当前题完成后停止。"
        job["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        return _progress_payload(job)


def get_evaluation_progress(job_id, mode=None):
    mode = _normalize_mode(mode) if mode else ""
    with EVALUATION_JOBS_LOCK:
        job = EVALUATION_JOBS.get(str(job_id or ""))
        if not job and mode:
            _expire_stale_evaluation_jobs(mode)
            job = _find_running_job(mode)
        if not job:
            return {
                "job_id": str(job_id or ""),
                "mode": mode,
                "status": "not_found",
                "total": 0,
                "completed": 0,
                "percent": 0,
                "current_case_id": "",
                "current_question": "",
                "elapsed_ms": 0,
                "error": "评测任务不存在或服务已重启。",
            }
        return _progress_payload(job)


def run_evaluation_for_admin(mode="deepseek"):
    mode = _normalize_mode(mode)
    previous = get_evaluation_snapshot(mode)
    try:
        return _execute_evaluation(mode)
    except Exception as exc:
        payload = dict(previous or {})
        payload.update({
            "ok": False,
            "ready": bool(previous and previous.get("ready")),
            "mode": mode,
            "error": str(exc),
            "message": "评测运行失败，已保留上一版缓存。" if previous and previous.get("ready") else str(exc),
            "updated_at": previous.get("updated_at", "") if previous else "",
        })
        return payload


def _join_items(items, empty="-"):
    values = _as_text_list(items)
    return "、".join(values) if values else empty


def _provider_summary_text(snapshot):
    stats = snapshot.get("provider_stats") or {}
    labels = [
        ("deepseek", "真实 DeepSeek"),
        ("direct_fact", "规则直答"),
        ("local", "本地回答"),
        ("local_fallback_after_llm_error", "本地兜底"),
        ("deepseek_error", "DeepSeek 失败"),
        ("unknown", "旧版未知来源"),
    ]
    parts = ["{0} {1} 题".format(label, stats.get(key)) for key, label in labels if stats.get(key)]
    return "；".join(parts) if parts else "暂无回答来源明细"


def _status_text(item):
    score = _score_value(item)
    if score >= 8:
        return "通过"
    if item.get("id") in HIGH_RISK_CASE_IDS:
        return "高风险"
    return "低分"


def _export_snapshots(mode):
    value = str(mode or "deepseek").strip().lower()
    modes = ["deepseek", "local"] if value == "combined" else [_normalize_mode(value)]
    snapshots = []
    for item_mode in modes:
        snapshot = get_evaluation_snapshot(item_mode)
        if snapshot.get("invalid_cache"):
            raise ValueError(snapshot.get("message") or DEEPSEEK_MIXED_CACHE_MESSAGE)
        if not snapshot.get("ready"):
            raise ValueError("请先运行 {0} 评测后再导出 Word 报告。".format("DeepSeek" if item_mode == "deepseek" else "本地基线"))
        if not snapshot.get("case_items"):
            raise ValueError("当前 {0} 评测缓存缺少测试用例明细，请重新运行评测后导出。".format("DeepSeek" if item_mode == "deepseek" else "本地基线"))
        snapshots.append(snapshot)
    return value, snapshots


def _add_snapshot_summary(document, snapshot):
    document.add_heading("{0} 评测摘要".format("本地基线" if snapshot.get("mode") == "local" else "DeepSeek"), level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    labels = [
        ("模式", "本地基线" if snapshot.get("mode") == "local" else "DeepSeek"),
        ("模型", snapshot.get("model", "-")),
        ("总分", "{0}%".format(snapshot.get("score_percent", 0))),
        ("事实准确率", "{0}%".format(round(float(snapshot.get("fact_accuracy", 0)) * 100, 2))),
        ("低分题数量", str(snapshot.get("failed_count", 0))),
        ("平均延迟", "{0}ms".format(snapshot.get("avg_latency_ms", 0))),
        ("回答来源", _provider_summary_text(snapshot)),
        ("最近更新时间", snapshot.get("updated_at", "-")),
    ]
    for index, (label, value) in enumerate(labels):
        if index:
            row = table.add_row()
        else:
            row = table.rows[0]
        row.cells[0].text = str(label)
        row.cells[1].text = str(value)


def _add_category_stats(document, snapshot):
    document.add_heading("分类通过率", level=2)
    stats = snapshot.get("category_stats") or _build_category_stats(snapshot.get("case_items") or [])
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["分类", "题数", "通过", "低分", "平均分"]
    for index, title in enumerate(headers):
        table.rows[0].cells[index].text = title
    for stat in stats:
        row = table.add_row().cells
        row[0].text = str(stat.get("category", ""))
        row[1].text = str(stat.get("total", 0))
        row[2].text = str(stat.get("passed", 0))
        row[3].text = str(stat.get("low_score", 0))
        row[4].text = str(stat.get("avg_score", 0))


def _add_low_score_items(document, snapshot):
    document.add_heading("低分题汇总", level=2)
    low_items = snapshot.get("low_score_items") or []
    if not low_items:
        document.add_paragraph("暂无低分题。")
        return
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["题号", "测试用例", "得分", "失败类型"]
    for index, title in enumerate(headers):
        table.rows[0].cells[index].text = title
    for item in low_items:
        row = table.add_row().cells
        row[0].text = str(item.get("id", ""))
        row[1].text = str(item.get("question", ""))
        row[2].text = str(item.get("score", 0))
        row[3].text = _join_items(item.get("failure_types"))


def _add_case_detail_table(document, snapshot):
    document.add_heading("完整测试用例明细", level=2)
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["题号", "分类", "测试用例", "预期结果", "实际结果", "得分", "命中情况", "失败类型"]
    for index, title in enumerate(headers):
        table.rows[0].cells[index].text = title
    for item in snapshot.get("case_items") or []:
        expected = item.get("expected") or {}
        hit = item.get("hit_detail") or {}
        expected_text = "应包含：{0}\n不应出现：{1}".format(
            _join_items(expected.get("must_include")),
            _join_items(expected.get("must_not_include")),
        )
        hit_text = "已命中：{0}\n未命中：{1}\n误触发：{2}".format(
            _join_items(hit.get("matched_include")),
            _join_items(hit.get("missing_include")),
            _join_items(hit.get("forbidden_hits")),
        )
        review = item.get("semantic_review") or {}
        if review.get("covered_include"):
            hit_text += "\n语义复核命中：{0}".format(_join_items(review.get("covered_include")))
        row = table.add_row().cells
        row[0].text = str(item.get("id", ""))
        row[1].text = str(item.get("category", ""))
        row[2].text = str(item.get("question", ""))
        row[3].text = expected_text
        row[4].text = str(item.get("answer", ""))
        row[5].text = "{0} 分\n{1}".format(item.get("score", 0), _status_text(item))
        row[6].text = hit_text
        row[7].text = _join_items(item.get("failure_types"))


def export_evaluation_docx(mode="deepseek"):
    from docx import Document

    export_mode, snapshots = _export_snapshots(mode)
    document = Document()
    document.add_heading("灵山胜境问答验收报告", 0)
    document.add_paragraph("报告范围：{0}".format("DeepSeek 与本地基线对比" if export_mode == "combined" else ("本地基线" if snapshots[0].get("mode") == "local" else "DeepSeek")))
    document.add_paragraph("导出时间：{0}".format(time.strftime("%Y-%m-%d %H:%M:%S")))
    document.add_paragraph("说明：预期结果采用 100 题评测数据集中的 must_include / must_not_include 验收条件。")
    for index, snapshot in enumerate(snapshots):
        if index:
            document.add_page_break()
        _add_snapshot_summary(document, snapshot)
        _add_category_stats(document, snapshot)
        _add_low_score_items(document, snapshot)
        _add_case_detail_table(document, snapshot)
    stream = BytesIO()
    document.save(stream)
    filename = "lingshan-qa-evaluation-{0}.docx".format(export_mode if export_mode == "combined" else snapshots[0].get("mode", "deepseek"))
    return stream.getvalue(), filename

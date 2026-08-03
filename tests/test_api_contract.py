# -*- coding: utf-8 -*-
from io import BytesIO
import json
import os
import sys
import tempfile
import unittest
from unittest.mock import patch

from docx import Document


ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BACKEND_DIR = os.path.join(ROOT_DIR, "backend")
if BACKEND_DIR not in sys.path:
    sys.path.insert(0, BACKEND_DIR)

import main  # noqa: E402
import evaluation_service  # noqa: E402
from asr_service import _clean_asr_text  # noqa: E402


class ApiContractTests(unittest.TestCase):
    def setUp(self):
        self.client = main.app.test_client()

    def admin_headers(self):
        with patch("auth_service.get_connection", side_effect=RuntimeError("mysql down")):
            with patch.dict(os.environ, {"AUTH_DEMO_FALLBACK": "1"}, clear=False):
                response = self.client.post(
                    "/api/auth/login",
                    json={"username": "admin", "password": "admin123456", "role": "admin"},
                )
        self.assertEqual(response.status_code, 200)
        return {"Authorization": "Bearer " + response.get_json()["token"]}

    def test_chat_response_contains_metadata_for_frontend_driving(self):
        response = self.client.post(
            "/api/chat",
            json={"query": "我带孩子玩四小时，请推荐路线", "history": [], "interest": "亲子家庭"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertIn("answer", data)
        self.assertIn("emotion", data)
        self.assertIn("route_suggestion", data)
        self.assertIn("sources", data)
        self.assertIn("latency_ms", data)
        self.assertEqual(data["route_suggestion"]["id"], "route_family")
        self.assertIn("recommendation_reason", data["route_suggestion"])
        self.assertIn("recommendation_context", data["route_suggestion"])

    def test_voice_response_contains_same_metadata_plus_tts_result(self):
        fake_tts = {
            "ok": True,
            "provider": "gpt_sovits",
            "engine": "gsv_tts_lite",
            "audio_url": "/audio/tts/fake.wav",
            "subtitles": [{"start_s": 0, "end_s": 1.2, "text": "九龙灌浴演出讲解"}],
        }
        with patch.object(main, "synthesize_speech", return_value=fake_tts):
            response = self.client.post(
                "/api/voice",
                json={"text": "九龙灌浴演出几点开始？", "history": [], "interest": "历史文化"},
            )

        self.assertIn(response.status_code, {200, 503})
        data = response.get_json()
        self.assertIn("answer", data)
        self.assertIn("emotion", data)
        self.assertIn("sources", data)
        self.assertIn("latency_ms", data)
        self.assertIn("tts", data)
        self.assertEqual(data["tts"]["engine"], "gsv_tts_lite")

    def test_uploaded_voice_asr_cleaning_applies_lingshan_hotword_corrections(self):
        self.assertEqual(_clean_asr_text("<|zh|>灵山饭宫和天下第一章"), "灵山梵宫和天下第一掌")

    def test_uploaded_voice_asr_cleaning_normalizes_noisy_realtime_phrases(self):
        cases = {
            "说一下介绍一下九龙观玉": "说一下介绍一下九龙灌浴",
            "要为介绍绍一阿玉王柱柱": "要我介绍一下阿育王柱",
            "百拜拜": "拜拜",
            "开开开喂喂喂听得见吗": "听得见吗",
        }
        for raw, expected in cases.items():
            with self.subTest(raw=raw):
                self.assertEqual(_clean_asr_text(raw), expected)

    def test_scenic_narration_returns_long_segmented_guide_text(self):
        response = self.client.get("/api/scenic/LS-006/narration")

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertIn("answer", data)
        self.assertGreaterEqual(len(data["answer"]), 350)
        self.assertLessEqual(len(data["answer"]), 650)
        self.assertIn("九龙灌浴", data["answer"])
        self.assertNotIn("菩提", data["answer"])
        self.assertIn("\n\n", data["answer"])
        self.assertNotIn("不是单独的拍照点", data["answer"])
        self.assertIn("display_segments", data)
        self.assertGreaterEqual(len(data["display_segments"]), 4)
        self.assertEqual(data["answer"], "\n\n".join(data["display_segments"]))
        self.assertIn("segments", data)
        self.assertGreaterEqual(len(data["segments"]), 4)
        self.assertLessEqual(len(data["segments"][0]), 32)
        self.assertTrue(all(30 <= len(segment) <= 90 for segment in data["segments"][1:]))
        self.assertIn(data["emotion"], {"neutral", "happy", "thanks", "surprised", "confused", "sad"})
        self.assertIn("sources", data)
        self.assertIn("latency_ms", data)

    def test_uploaded_voice_realtime_uses_deepseek_corrected_asr_text_for_dialogue(self):
        fake_tts = {
            "ok": True,
            "audio_url": "/audio/tts/fake.wav",
            "engine": "gsv_tts_lite",
        }
        captured = {}

        def fake_chat_detail(query, history=None, interest=None):
            captured["query"] = query
            return {
                "answer": "九龙灌浴常见场次为 10:00、11:30、13:30、15:00。",
                "emotion": "neutral",
                "emotion_label": "自然",
                "route_suggestion": None,
                "sources": [],
                "latency_ms": 12,
            }

        with patch.object(main, "save_uploaded_audio", return_value="voice.webm"):
            with patch.object(main, "transcribe_audio", return_value={"ok": True, "text": "九龙观浴几点开始", "provider": "SenseVoice"}):
                with patch.object(
                    main,
                    "correct_asr_text",
                    return_value={
                        "text": "九龙灌浴几点开始",
                        "original_text": "九龙观浴几点开始",
                        "corrected_text": "九龙灌浴几点开始",
                        "pre_llm_text": "九龙观浴几点开始",
                        "leading_noise_removed": False,
                        "leading_noise_reason": "",
                        "llm_corrected": True,
                        "correction_provider": "deepseek",
                        "correction_confidence": 0.91,
                        "correction_error": "",
                    },
                ):
                    with patch.object(main.rag, "chat_detail", side_effect=fake_chat_detail):
                        with patch.object(main, "synthesize_speech", return_value=fake_tts):
                            response = self.client.post(
                                "/api/voice/realtime",
                                data={"audio": (BytesIO(b"fake-audio"), "voice.webm")},
                                content_type="multipart/form-data",
                            )

        self.assertEqual(200, response.status_code)
        data = response.get_json()
        self.assertEqual("九龙灌浴几点开始", data["query"])
        self.assertTrue(data["asr"]["llm_corrected"])
        self.assertEqual("九龙观浴几点开始", data["asr"]["original_text"])
        self.assertEqual("九龙观浴几点开始", data["asr"]["pre_llm_text"])
        self.assertFalse(data["asr"]["leading_noise_removed"])
        self.assertIn("九龙灌浴几点开始", captured["query"])

    def test_uploaded_voice_realtime_criticism_uses_sad_emotion_from_corrected_query(self):
        fake_tts = {
            "ok": True,
            "audio_url": "/audio/tts/fake.wav",
            "engine": "gsv_tts_lite",
        }

        def fake_chat_detail(query, history=None, interest=None):
            return {
                "answer": "抱歉，刚才没有讲清楚。我重新为您讲解灵山大照壁。",
                "emotion": "neutral",
                "emotion_label": "自然",
                "route_suggestion": None,
                "sources": [],
                "latency_ms": 12,
            }

        with patch.object(main, "save_uploaded_audio", return_value="voice.webm"):
            with patch.object(main, "transcribe_audio", return_value={"ok": True, "text": "可能头来了讲的什么东西啊啊太烂了", "provider": "SenseVoice"}):
                with patch.object(
                    main,
                    "correct_asr_text",
                    return_value={
                        "text": "讲的什么东西啊太烂了",
                        "original_text": "可能头来了讲的什么东西啊啊太烂了",
                        "corrected_text": "讲的什么东西啊太烂了",
                        "pre_llm_text": "讲的什么东西啊啊太烂了",
                        "leading_noise_removed": True,
                        "leading_noise_reason": "移除意图前无意义识别片段",
                        "llm_corrected": True,
                        "correction_provider": "deepseek",
                        "correction_confidence": 0.9,
                        "correction_error": "",
                    },
                ):
                    with patch.object(main.rag, "chat_detail", side_effect=fake_chat_detail):
                        with patch.object(main, "synthesize_speech", return_value=fake_tts):
                            response = self.client.post(
                                "/api/voice/realtime",
                                data={"audio": (BytesIO(b"fake-audio"), "voice.webm")},
                                content_type="multipart/form-data",
                            )

        self.assertEqual(200, response.status_code)
        data = response.get_json()
        self.assertEqual("讲的什么东西啊太烂了", data["query"])
        self.assertEqual("sad", data["emotion"])
        self.assertEqual("伤心反思", data["emotion_label"])

    def test_uploaded_voice_realtime_stops_when_deepseek_correction_fails(self):
        correction = {
            "ok": True,
            "text": "可能千万万讲的什么东西啊太烂了",
            "original_text": "可能千万万讲的什么东西啊太烂了",
            "corrected_text": "可能千万万讲的什么东西啊太烂了",
            "pre_llm_text": "可能千万万讲的什么东西啊太烂了",
            "leading_noise_removed": False,
            "leading_noise_reason": "",
            "llm_corrected": False,
            "correction_provider": "deepseek",
            "correction_confidence": 0.0,
            "correction_reason": "",
            "correction_error": "DeepSeek 超时",
        }

        with patch.object(main, "save_uploaded_audio", return_value="voice.webm"):
            with patch.object(main, "transcribe_audio", return_value={"ok": True, "text": "可能千万万讲的什么东西啊太烂了", "provider": "SenseVoice"}):
                with patch.object(main, "correct_asr_text", return_value=correction):
                    with patch.object(main.rag, "chat_detail", side_effect=AssertionError("纠错失败时不应进入对话")):
                        response = self.client.post(
                            "/api/voice/realtime",
                            data={"audio": (BytesIO(b"fake-audio"), "voice.webm")},
                            content_type="multipart/form-data",
                        )

        self.assertEqual(503, response.status_code)
        data = response.get_json()
        self.assertFalse(data["ok"])
        self.assertTrue(data["asr"]["correction_failed"])
        self.assertIn("语音纠错失败", data["error"])
        self.assertIn("DeepSeek 超时", data["asr"]["correction_error"])

    def test_voice_correct_text_endpoint_uses_deepseek_asr_correction_metadata(self):
        correction = {
            "ok": True,
            "text": "讲的什么东西啊太烂了",
            "original_text": "可能千万万讲的什么东西啊太烂了",
            "corrected_text": "讲的什么东西啊太烂了",
            "pre_llm_text": "可能千万万讲的什么东西啊太烂了",
            "leading_noise_removed": False,
            "leading_noise_reason": "",
            "llm_corrected": True,
            "correction_provider": "deepseek",
            "correction_confidence": 0.93,
            "correction_reason": "删除开头无意义误识别片段",
            "correction_error": "",
        }

        with patch.object(main, "correct_asr_text", return_value=correction) as corrected:
            response = self.client.post(
                "/api/voice/correct-text",
                json={
                    "text": "可能千万万讲的什么东西啊太烂了",
                    "history": [{"role": "assistant", "content": "上一轮讲解"}],
                    "realtime": False,
                },
            )

        self.assertEqual(200, response.status_code)
        data = response.get_json()
        self.assertEqual("讲的什么东西啊太烂了", data["text"])
        self.assertTrue(data["llm_corrected"])
        self.assertEqual("deepseek", data["correction_provider"])
        corrected.assert_called_once()
        self.assertEqual("可能千万万讲的什么东西啊太烂了", corrected.call_args[0][0])
        self.assertEqual([{"role": "assistant", "content": "上一轮讲解"}], corrected.call_args[1]["history"])

    def test_voice_correct_text_endpoint_reports_correction_failure(self):
        correction = {
            "ok": True,
            "text": "可能千万万讲的什么东西啊太烂了",
            "original_text": "可能千万万讲的什么东西啊太烂了",
            "corrected_text": "可能千万万讲的什么东西啊太烂了",
            "pre_llm_text": "可能千万万讲的什么东西啊太烂了",
            "leading_noise_removed": False,
            "leading_noise_reason": "",
            "llm_corrected": False,
            "correction_provider": "deepseek",
            "correction_confidence": 0.0,
            "correction_reason": "",
            "correction_error": "DeepSeek 超时",
        }

        with patch.object(main, "correct_asr_text", return_value=correction):
            response = self.client.post(
                "/api/voice/correct-text",
                json={"text": "可能千万万讲的什么东西啊太烂了"},
            )

        self.assertEqual(503, response.status_code)
        data = response.get_json()
        self.assertFalse(data["ok"])
        self.assertTrue(data["correction_failed"])
        self.assertIn("DeepSeek 超时", data["correction_error"])

    def test_chat_criticism_response_uses_sad_reflective_emotion_label(self):
        response = self.client.post(
            "/api/chat",
            json={"query": "可能千万万讲的什么东西啊太烂了", "history": [], "interest": ""},
        )

        self.assertEqual(200, response.status_code)
        data = response.get_json()
        self.assertEqual("sad", data["emotion"])
        self.assertEqual("伤心反思", data["emotion_label"])

    def test_scenic_narration_uses_spot_specific_natural_copy(self):
        wall = self.client.get("/api/scenic/LS-001/narration").get_json()
        bath = self.client.get("/api/scenic/LS-006/narration").get_json()

        self.assertNotEqual(wall["display_segments"][1], bath["display_segments"][1])
        self.assertNotIn("先在这里停一下，听我把重点带起来", wall["answer"])
        self.assertNotIn("先在这里停一下，听我把重点带起来", bath["answer"])
        self.assertNotIn("在整条游线中的位置很讲究", bath["answer"])
        self.assertNotIn("入口、广场、建筑和后面的核心景观串起来", bath["answer"])
        self.assertNotIn("看见什么", wall["answer"])
        self.assertNotIn("看见什么", bath["answer"])
        self.assertNotIn("这处景观更像一段开场白", wall["answer"])
        self.assertNotIn("这处景观更像一段开场白", bath["answer"])
        self.assertNotIn("把刚才看到的形象、声音和空间感连起来理解", wall["answer"])
        self.assertNotIn("把刚才看到的形象、声音和空间感连起来理解", bath["answer"])
        self.assertIn("灵山大照壁", wall["answer"])
        self.assertIn("九龙灌浴", bath["answer"])

    def test_scenic_narration_keeps_first_voice_segment_short_for_fast_start(self):
        for scenic_id in ["LS-001", "LS-006", "LS-011"]:
            with self.subTest(scenic_id=scenic_id):
                data = self.client.get("/api/scenic/{0}/narration".format(scenic_id)).get_json()

                self.assertLessEqual(len(data["segments"][0]), 32)
                self.assertGreaterEqual(len(data["segments"]), 4)

    def test_scenic_narration_unknown_spot_returns_404(self):
        response = self.client.get("/api/scenic/UNKNOWN/narration")

        self.assertEqual(response.status_code, 404)
        self.assertEqual(response.get_json()["error"], "景点不存在")

    def test_admin_evaluation_snapshot_and_run_endpoints_are_available(self):
        headers = self.admin_headers()
        cached = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "score_percent": 96.8,
            "fact_accuracy": 0.94,
            "failed_count": 3,
            "avg_latency_ms": 812.4,
            "updated_at": "2026-07-11 10:00:00",
            "low_score_items": [{"id": "Q022", "question": "《吉祥颂》每天什么时间演出？", "score": 6.5}],
            "high_risk_failures": [],
            "category_stats": [{"category": "performance", "total": 1, "passed": 0, "avg_score": 6.5}],
            "case_items": [
                {
                    "id": "Q022",
                    "category": "performance",
                    "question": "《吉祥颂》每天什么时间演出？",
                    "expected": {"must_include": ["吉祥颂", "10:35"], "must_not_include": ["九龙灌浴"], "weight": 1},
                    "answer": "九龙灌浴常见场次为 10:00、11:30、13:30、15:00。",
                    "score": 3.0,
                    "hit_detail": {
                        "matched_include": ["吉祥颂"],
                        "missing_include": ["10:35"],
                        "forbidden_hits": ["九龙灌浴"],
                    },
                    "semantic_review": {
                        "reviewed": True,
                        "adjusted": True,
                        "covered_include": ["吉祥颂"],
                        "evidence": "回答中提到相关演出名。",
                        "confidence": 0.89,
                        "error": "",
                    },
                    "failure_types": ["intent_misroute"],
                    "latency_ms": 812,
                }
            ],
        }
        with patch.object(main, "get_evaluation_snapshot", return_value=cached):
            response = self.client.get("/api/admin/evaluation?mode=deepseek", headers=headers)

        self.assertEqual(response.status_code, 200)
        evaluation = response.get_json()["evaluation"]
        self.assertEqual(evaluation["mode"], "deepseek")
        self.assertEqual(evaluation["model"], "deepseek-chat")
        self.assertIn("case_items", evaluation)
        self.assertEqual("Q022", evaluation["case_items"][0]["id"])
        self.assertEqual(["吉祥颂", "10:35"], evaluation["case_items"][0]["expected"]["must_include"])
        self.assertEqual(["九龙灌浴"], evaluation["case_items"][0]["hit_detail"]["forbidden_hits"])
        self.assertIn("category_stats", evaluation)

        progress = {
            "job_id": "eval-deepseek-unit",
            "mode": "deepseek",
            "status": "running",
            "total": 100,
            "completed": 0,
            "percent": 0,
            "current_case_id": "",
            "current_question": "",
            "elapsed_ms": 0,
            "error": "",
        }
        with patch.object(main, "start_evaluation_job", return_value={"job_id": "eval-deepseek-unit", "progress": progress}):
            response = self.client.post(
                "/api/admin/evaluation/run",
                json={"mode": "deepseek"},
                headers=headers,
            )

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["job_id"], "eval-deepseek-unit")
        self.assertEqual(payload["progress"]["status"], "running")
        self.assertEqual(payload["progress"]["percent"], 0)

        refreshed = dict(cached, score_percent=98.1, updated_at="2026-07-11 10:05:00")
        completed = dict(progress, status="completed", completed=100, percent=100, evaluation=refreshed)
        with patch.object(main, "get_evaluation_progress", return_value=completed):
            response = self.client.get(
                "/api/admin/evaluation/progress?job_id=eval-deepseek-unit",
                headers=headers,
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.get_json()["progress"]["status"], "completed")
        self.assertEqual(response.get_json()["progress"]["evaluation"]["score_percent"], 98.1)

    def test_evaluation_job_does_not_reuse_stale_waiting_progress(self):
        class DeadThread(object):
            def is_alive(self):
                return False

        class FakeThread(object):
            def __init__(self, target=None, args=(), daemon=False):
                self.target = target
                self.args = args
                self.daemon = daemon
                self.started = False

            def start(self):
                self.started = True

            def is_alive(self):
                return self.started

        with evaluation_service.EVALUATION_JOBS_LOCK:
            evaluation_service.EVALUATION_JOBS.clear()
            evaluation_service.EVALUATION_JOBS["eval-deepseek-stale"] = {
                "job_id": "eval-deepseek-stale",
                "mode": "deepseek",
                "status": "queued",
                "total": 100,
                "completed": 0,
                "percent": 0,
                "current_case_id": "",
                "current_question": "",
                "message": "评测任务已创建",
                "error": "",
                "started_at": 1,
                "updated_at": "2026-01-01 00:00:00",
                "thread": DeadThread(),
            }

        try:
            with patch.object(evaluation_service.threading, "Thread", FakeThread):
                payload = evaluation_service.start_evaluation_job("deepseek")
            self.assertNotEqual("eval-deepseek-stale", payload["job_id"])
            self.assertEqual("running", payload["progress"]["status"])
            self.assertEqual(100, payload["progress"]["total"])
            self.assertIn("正在启动", payload["progress"]["message"])
            with evaluation_service.EVALUATION_JOBS_LOCK:
                stale = evaluation_service.EVALUATION_JOBS["eval-deepseek-stale"]
            self.assertEqual("failed", stale["status"])
            self.assertIn("已停止响应", stale["error"])
        finally:
            with evaluation_service.EVALUATION_JOBS_LOCK:
                evaluation_service.EVALUATION_JOBS.clear()

    def test_evaluation_job_accepts_client_job_id_for_timeout_recovery(self):
        class FakeThread(object):
            def __init__(self, target=None, args=(), daemon=False):
                self.target = target
                self.args = args
                self.daemon = daemon
                self.started = False

            def start(self):
                self.started = True

            def is_alive(self):
                return self.started

        client_job_id = "eval-deepseek-client-unit-001"
        try:
            with evaluation_service.EVALUATION_JOBS_LOCK:
                evaluation_service.EVALUATION_JOBS.clear()
            with patch.object(evaluation_service.threading, "Thread", FakeThread):
                payload = evaluation_service.start_evaluation_job("deepseek", job_id=client_job_id)
            self.assertEqual(client_job_id, payload["job_id"])
            self.assertEqual("running", payload["progress"]["status"])

            progress = evaluation_service.get_evaluation_progress(client_job_id, mode="deepseek")
            self.assertEqual(client_job_id, progress["job_id"])
            self.assertEqual("running", progress["status"])

            fallback = evaluation_service.get_evaluation_progress("missing-client-job", mode="deepseek")
            self.assertEqual(client_job_id, fallback["job_id"])
            self.assertEqual("running", fallback["status"])
        finally:
            with evaluation_service.EVALUATION_JOBS_LOCK:
                evaluation_service.EVALUATION_JOBS.clear()

    def test_evaluation_job_can_be_cancelled_without_overwriting_cache(self):
        class AliveThread(object):
            def is_alive(self):
                return True

        job_id = "eval-deepseek-client-cancel-001"
        try:
            with evaluation_service.EVALUATION_JOBS_LOCK:
                evaluation_service.EVALUATION_JOBS.clear()
                evaluation_service.EVALUATION_JOBS[job_id] = {
                    "job_id": job_id,
                    "mode": "deepseek",
                    "status": "running",
                    "total": 100,
                    "completed": 9,
                    "percent": 9,
                    "current_case_id": "Q010",
                    "current_question": "测试取消",
                    "message": "评测进行中",
                    "error": "",
                    "started_at": 1,
                    "updated_at": "2026-01-01 00:00:00",
                    "thread": AliveThread(),
                }

            payload = evaluation_service.cancel_evaluation_job("deepseek", job_id)

            self.assertEqual(job_id, payload["job_id"])
            self.assertEqual("cancelling", payload["status"])
            self.assertTrue(payload["cancel_requested"])
            self.assertIn("正在结束", payload["message"])

            progress = evaluation_service.get_evaluation_progress(job_id, mode="deepseek")
            self.assertEqual("cancelling", progress["status"])
            self.assertTrue(progress["cancel_requested"])

            missing = evaluation_service.cancel_evaluation_job("deepseek", "missing-job")
            self.assertEqual("not_found", missing["status"])
        finally:
            with evaluation_service.EVALUATION_JOBS_LOCK:
                evaluation_service.EVALUATION_JOBS.clear()

    def test_admin_evaluation_cancel_endpoint_returns_progress(self):
        headers = self.admin_headers()
        progress = {
            "job_id": "eval-deepseek-client-cancel-002",
            "mode": "deepseek",
            "status": "cancelling",
            "total": 100,
            "completed": 12,
            "percent": 12,
            "current_case_id": "Q013",
            "current_question": "正在结束",
            "elapsed_ms": 1200,
            "message": "正在结束评测，当前题完成后停止。",
            "error": "",
            "cancel_requested": True,
        }
        with patch.object(main, "cancel_evaluation_job", return_value=progress):
            response = self.client.post(
                "/api/admin/evaluation/cancel",
                json={"mode": "deepseek", "job_id": "eval-deepseek-client-cancel-002"},
                headers=headers,
            )

        self.assertEqual(200, response.status_code)
        payload = response.get_json()
        self.assertEqual("cancelling", payload["progress"]["status"])
        self.assertTrue(payload["progress"]["cancel_requested"])

    def test_admin_evaluation_review_low_score_endpoint_starts_deepseek_job_only(self):
        headers = self.admin_headers()
        progress = {
            "job_id": "review-deepseek-unit",
            "mode": "deepseek",
            "status": "running",
            "total": 2,
            "completed": 0,
            "percent": 0,
            "current_case_id": "",
            "current_question": "正在启动低分题语义复核",
            "elapsed_ms": 0,
            "message": "正在启动低分题语义复核",
            "error": "",
        }
        with patch.object(
            main,
            "start_semantic_review_job",
            return_value={"job_id": "review-deepseek-unit", "progress": progress},
            create=True,
        ):
            response = self.client.post(
                "/api/admin/evaluation/review-low-score",
                json={"mode": "deepseek", "job_id": "review-deepseek-unit"},
                headers=headers,
            )

        self.assertEqual(200, response.status_code)
        payload = response.get_json()
        self.assertEqual("review-deepseek-unit", payload["job_id"])
        self.assertEqual("running", payload["progress"]["status"])

        response = self.client.post(
            "/api/admin/evaluation/review-low-score",
            json={"mode": "local"},
            headers=headers,
        )

        self.assertEqual(400, response.status_code)
        self.assertIn("仅支持 DeepSeek", response.get_json()["error"])

    def test_semantic_review_existing_deepseek_cache_rewrites_summary_without_rerunning_cases(self):
        legacy_cache = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "scoring_version": "relaxed_keyword_v3",
            "case_count": 1,
            "score_percent": 67.5,
            "fact_accuracy": 0.0,
            "failed_count": 1,
            "avg_latency_ms": 2100,
            "updated_at": "2026-07-14 12:10:00",
            "provider_stats": {"deepseek": 1},
            "case_items": [
                {
                    "id": "Q010",
                    "category": "culture",
                    "question": "灵山梵宫的核心价值是什么？",
                    "expected": {
                        "must_include": ["世界佛教论坛永久会址"],
                        "must_not_include": ["拈花湾"],
                        "source_doc": "unit-test",
                        "weight": 1,
                    },
                    "answer": "核心在于它是世界佛教论坛的永久会址。",
                    "answer_provider": "deepseek",
                    "score": 6.75,
                    "hit_detail": {
                        "matched_include": [],
                        "missing_include": ["世界佛教论坛永久会址"],
                        "forbidden_hits": [],
                    },
                    "failure_types": ["missing_fact"],
                    "latency_ms": 2100,
                    "sources": [],
                }
            ],
        }

        class FakeEvaluator(object):
            SEMANTIC_REVIEW_VERSION = "llm_low_score_review_v1"

            def semantic_review_low_score_items(self, items, reviewer=None, progress_callback=None, cancel_checker=None, mode="deepseek"):
                adjusted = [dict(items[0])]
                adjusted[0]["score"] = 10.0
                adjusted[0]["hit_detail"] = {
                    "include_hits": 1,
                    "include_total": 1,
                    "matched_include": ["世界佛教论坛永久会址"],
                    "missing_include": [],
                    "forbidden_hits": [],
                }
                adjusted[0]["failure_types"] = []
                adjusted[0]["semantic_review"] = {
                    "reviewed": True,
                    "adjusted": True,
                    "covered_include": ["世界佛教论坛永久会址"],
                    "evidence": "语义等价。",
                    "confidence": 0.93,
                    "error": "",
                }
                return adjusted, {"reviewed": 1, "adjusted": 1, "errors": 0}

        with tempfile.TemporaryDirectory() as cache_dir:
            summary_path = os.path.join(cache_dir, "lingshan_eval_deepseek.json")
            latest_path = os.path.join(cache_dir, "lingshan_eval_deepseek_latest.json")
            with open(summary_path, "w", encoding="utf-8") as f:
                json.dump(legacy_cache, f, ensure_ascii=False)

            with patch.object(evaluation_service, "CACHE_DIR", cache_dir), patch.object(
                evaluation_service, "_load_evaluator", return_value=FakeEvaluator()
            ):
                payload = evaluation_service._execute_semantic_review("deepseek")
                with open(summary_path, "r", encoding="utf-8") as f:
                    persisted = json.load(f)
                with open(latest_path, "r", encoding="utf-8") as f:
                    latest = json.load(f)

        self.assertEqual("llm_low_score_review_v1", payload["semantic_review_version"])
        self.assertEqual(100.0, payload["score_percent"])
        self.assertEqual(1.0, payload["fact_accuracy"])
        self.assertEqual(0, payload["failed_count"])
        self.assertEqual(["世界佛教论坛永久会址"], payload["case_items"][0]["semantic_review"]["covered_include"])
        self.assertEqual("llm_low_score_review_v1", persisted["semantic_review_version"])
        self.assertEqual("Q010", latest["case_items"][0]["id"])
        with self.assertRaisesRegex(ValueError, "仅支持 DeepSeek"):
            evaluation_service._execute_semantic_review("local")

    def test_admin_evaluation_export_returns_word_document(self):
        headers = self.admin_headers()
        doc = Document()
        doc.add_heading("灵山胜境问答验收报告", 0)
        doc.add_paragraph("测试用例")
        doc.add_paragraph("预期结果")
        doc.add_paragraph("实际结果")
        stream = BytesIO()
        doc.save(stream)

        with patch.object(
            main,
            "export_evaluation_docx",
            return_value=(stream.getvalue(), "lingshan-qa-evaluation-deepseek.docx"),
            create=True,
        ):
            response = self.client.get("/api/admin/evaluation/export?mode=deepseek", headers=headers)

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            response.content_type,
        )
        self.assertIn("lingshan-qa-evaluation-deepseek.docx", response.headers["Content-Disposition"])
        exported = Document(BytesIO(response.data))
        text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        self.assertIn("测试用例", text)
        self.assertIn("预期结果", text)
        self.assertIn("实际结果", text)

    def test_evaluation_export_service_generates_case_detail_word_document(self):
        snapshot = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "score_percent": 96.8,
            "fact_accuracy": 0.94,
            "failed_count": 1,
            "avg_latency_ms": 812.4,
            "updated_at": "2026-07-11 10:00:00",
            "provider_stats": {"deepseek": 1},
            "category_stats": [{"category": "performance", "total": 1, "passed": 0, "low_score": 1, "avg_score": 3.0}],
            "low_score_items": [{"id": "Q022", "question": "《吉祥颂》每天什么时间演出？", "score": 3.0}],
            "case_items": [
                {
                    "id": "Q022",
                    "category": "performance",
                    "question": "《吉祥颂》每天什么时间演出？",
                    "expected": {"must_include": ["吉祥颂", "10:35"], "must_not_include": ["九龙灌浴"], "weight": 1},
                    "answer": "九龙灌浴常见场次为 10:00、11:30、13:30、15:00。",
                    "score": 3.0,
                    "hit_detail": {
                        "matched_include": ["吉祥颂"],
                        "missing_include": ["10:35"],
                        "forbidden_hits": ["九龙灌浴"],
                    },
                    "semantic_review": {
                        "reviewed": True,
                        "adjusted": True,
                        "covered_include": ["吉祥颂"],
                        "evidence": "回答中提到相关演出名。",
                        "confidence": 0.89,
                        "error": "",
                    },
                    "failure_types": ["intent_misroute"],
                    "latency_ms": 812,
                }
            ],
        }

        with patch.object(evaluation_service, "get_evaluation_snapshot", return_value=snapshot):
            content, filename = evaluation_service.export_evaluation_docx("deepseek")

        self.assertEqual(filename, "lingshan-qa-evaluation-deepseek.docx")
        exported = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        table_text = "\n".join(cell.text for table in exported.tables for row in table.rows for cell in row.cells)
        self.assertIn("灵山胜境问答验收报告", text)
        self.assertIn("回答来源", table_text)
        self.assertIn("真实 DeepSeek 1 题", table_text)
        self.assertIn("测试用例", table_text)
        self.assertIn("预期结果", table_text)
        self.assertIn("实际结果", table_text)
        self.assertIn("语义复核命中", table_text)
        self.assertIn("《吉祥颂》每天什么时间演出？", table_text)

    def test_evaluation_export_service_requires_ready_case_details(self):
        with patch.object(evaluation_service, "get_evaluation_snapshot", return_value={"ready": False, "mode": "deepseek"}):
            with self.assertRaisesRegex(ValueError, "请先运行 DeepSeek 评测"):
                evaluation_service.export_evaluation_docx("deepseek")

    def test_deepseek_mixed_provider_cache_is_marked_invalid_and_not_exportable(self):
        mixed_cache = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "case_count": 2,
            "score_percent": 99.0,
            "fact_accuracy": 1.0,
            "failed_count": 0,
            "avg_latency_ms": 3000,
            "updated_at": "2026-07-14 10:10:00",
            "provider_stats": {"deepseek": 1, "direct_fact": 1},
            "case_items": [
                {
                    "id": "Q001",
                    "category": "overview",
                    "question": "灵山胜境在哪里？",
                    "expected": {"must_include": ["无锡"], "must_not_include": ["拈花湾"], "weight": 1},
                    "answer": "灵山胜境位于无锡。",
                    "answer_provider": "deepseek",
                    "score": 10,
                    "hit_detail": {"matched_include": ["无锡"], "missing_include": [], "forbidden_hits": []},
                    "failure_types": [],
                    "latency_ms": 6000,
                },
                {
                    "id": "Q002",
                    "category": "scenic",
                    "question": "灵山大佛有多高？",
                    "expected": {"must_include": ["88米"], "must_not_include": ["拈花湾"], "weight": 1},
                    "answer": "灵山大佛高88米。",
                    "answer_provider": "direct_fact",
                    "score": 10,
                    "hit_detail": {"matched_include": ["88米"], "missing_include": [], "forbidden_hits": []},
                    "failure_types": [],
                    "latency_ms": 1,
                },
            ],
        }

        with tempfile.TemporaryDirectory() as cache_dir:
            summary_path = os.path.join(cache_dir, "lingshan_eval_deepseek.json")
            with open(summary_path, "w", encoding="utf-8") as f:
                json.dump(mixed_cache, f, ensure_ascii=False)

            with patch.object(evaluation_service, "CACHE_DIR", cache_dir):
                snapshot = evaluation_service.get_evaluation_snapshot("deepseek")
                with self.assertRaisesRegex(ValueError, "旧缓存包含规则直答"):
                    evaluation_service.export_evaluation_docx("deepseek")

        self.assertFalse(snapshot["ok"])
        self.assertFalse(snapshot["ready"])
        self.assertTrue(snapshot["invalid_cache"])
        self.assertIn("旧缓存包含规则直答", snapshot["message"])

    def test_evaluation_export_recovers_legacy_summary_cache_from_latest_full_result(self):
        legacy_summary = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "case_count": 1,
            "score_percent": 90.0,
            "fact_accuracy": 1.0,
            "failed_count": 0,
            "avg_latency_ms": 10,
            "updated_at": "2026-07-14 09:50:07",
            "low_score_items": [],
            "high_risk_failures": [],
        }
        latest_full = dict(
            legacy_summary,
            items=[
                {
                    "id": "Q001",
                    "category": "overview",
                    "question": "灵山胜境在哪里？",
                    "expected": {
                        "must_include": ["无锡"],
                        "must_not_include": ["拈花湾"],
                        "source_doc": "unit-test",
                        "weight": 1,
                    },
                    "answer": "灵山胜境位于无锡。",
                    "answer_provider": "deepseek",
                    "score": 10,
                    "hit_detail": {
                        "matched_include": ["无锡"],
                        "missing_include": [],
                        "forbidden_hits": [],
                    },
                    "failure_types": [],
                    "latency_ms": 10,
                    "sources": [],
                }
            ],
        )

        with tempfile.TemporaryDirectory() as cache_dir:
            summary_path = os.path.join(cache_dir, "lingshan_eval_deepseek.json")
            latest_path = os.path.join(cache_dir, "lingshan_eval_deepseek_latest.json")
            with open(summary_path, "w", encoding="utf-8") as f:
                json.dump(legacy_summary, f, ensure_ascii=False)
            with open(latest_path, "w", encoding="utf-8") as f:
                json.dump(latest_full, f, ensure_ascii=False)

            with patch.object(evaluation_service, "CACHE_DIR", cache_dir):
                content, filename = evaluation_service.export_evaluation_docx("deepseek")
                with open(summary_path, "r", encoding="utf-8") as f:
                    migrated = json.load(f)

        self.assertEqual(filename, "lingshan-qa-evaluation-deepseek.docx")
        self.assertGreater(len(content), 1000)
        self.assertIn("case_items", migrated)
        self.assertEqual("Q001", migrated["case_items"][0]["id"])

    def test_evaluation_snapshot_rescores_legacy_keyword_matches_without_rerun(self):
        legacy_cache = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "case_count": 2,
            "score_percent": 51.25,
            "fact_accuracy": 0.0,
            "failed_count": 2,
            "avg_latency_ms": 6400,
            "updated_at": "2026-07-14 12:00:00",
            "provider_stats": {"deepseek": 2},
            "low_score_items": [],
            "high_risk_failures": [],
            "case_items": [
                {
                    "id": "Q013",
                    "category": "ticket",
                    "question": "观光车联票多少钱？",
                    "expected": {"must_include": ["225元", "门票+观光车"], "must_not_include": ["拈花湾"], "weight": 1},
                    "answer": "观光车联票的价格是225元。这个联票包含了景区门票和观光车。",
                    "answer_provider": "deepseek",
                    "score": 6.75,
                    "hit_detail": {"matched_include": ["225元"], "missing_include": ["门票+观光车"], "forbidden_hits": []},
                    "failure_types": ["missing_fact"],
                    "latency_ms": 6476,
                },
                {
                    "id": "Q015",
                    "category": "visit_time",
                    "question": "灵山胜境最佳游览季节是什么时候？",
                    "expected": {"must_include": ["3-5月", "9-11月"], "must_not_include": ["拈花湾"], "weight": 1},
                    "answer": "来灵山胜境，我最推荐的季节是春秋两季，也就是每年的3月到5月，以及9月到11月。",
                    "answer_provider": "deepseek",
                    "score": 3.5,
                    "hit_detail": {"matched_include": [], "missing_include": ["3-5月", "9-11月"], "forbidden_hits": []},
                    "failure_types": ["missing_fact"],
                    "latency_ms": 6385,
                },
            ],
        }

        with tempfile.TemporaryDirectory() as cache_dir:
            summary_path = os.path.join(cache_dir, "lingshan_eval_deepseek.json")
            with open(summary_path, "w", encoding="utf-8") as f:
                json.dump(legacy_cache, f, ensure_ascii=False)

            with patch.object(evaluation_service, "CACHE_DIR", cache_dir):
                snapshot = evaluation_service.get_evaluation_snapshot("deepseek")
                with open(summary_path, "r", encoding="utf-8") as f:
                    persisted = json.load(f)

        self.assertEqual("relaxed_keyword_v3", snapshot["scoring_version"])
        self.assertEqual(100.0, snapshot["score_percent"])
        self.assertEqual(1.0, snapshot["fact_accuracy"])
        self.assertEqual(0, snapshot["failed_count"])
        self.assertEqual([], snapshot["low_score_items"])
        self.assertEqual(["225元", "门票+观光车"], snapshot["case_items"][0]["hit_detail"]["matched_include"])
        self.assertEqual([], snapshot["case_items"][0]["hit_detail"]["missing_include"])
        self.assertEqual(["3-5月", "9-11月"], snapshot["case_items"][1]["hit_detail"]["matched_include"])
        self.assertEqual([], snapshot["case_items"][1]["hit_detail"]["missing_include"])
        self.assertEqual("relaxed_keyword_v3", persisted["scoring_version"])

    def test_evaluation_snapshot_rescores_v2_year_range_cache_without_rerun(self):
        legacy_cache = {
            "ok": True,
            "ready": True,
            "mode": "deepseek",
            "model": "deepseek-chat",
            "scoring_version": "relaxed_keyword_v2",
            "case_count": 1,
            "score_percent": 78.3,
            "fact_accuracy": 0.0,
            "failed_count": 1,
            "avg_latency_ms": 6496,
            "updated_at": "2026-07-14 20:44:05",
            "provider_stats": {"deepseek": 1},
            "case_items": [
                {
                    "id": "Q007",
                    "category": "history",
                    "question": "祥符禅寺什么时候得名？",
                    "expected": {
                        "must_include": ["北宋", "大中祥符", "1008-1016"],
                        "must_not_include": ["拈花湾"],
                        "source_doc": "unit-test",
                        "weight": 1,
                    },
                    "answer": "祥符禅寺得名于北宋大中祥符年间，具体时间约为公元1008年至1016年。",
                    "answer_provider": "deepseek",
                    "score": 7.83,
                    "hit_detail": {
                        "matched_include": ["北宋", "大中祥符"],
                        "missing_include": ["1008-1016"],
                        "forbidden_hits": [],
                    },
                    "failure_types": ["missing_fact"],
                    "latency_ms": 6496,
                }
            ],
        }

        with tempfile.TemporaryDirectory() as cache_dir:
            summary_path = os.path.join(cache_dir, "lingshan_eval_deepseek.json")
            with open(summary_path, "w", encoding="utf-8") as f:
                json.dump(legacy_cache, f, ensure_ascii=False)

            with patch.object(evaluation_service, "CACHE_DIR", cache_dir):
                snapshot = evaluation_service.get_evaluation_snapshot("deepseek")
                with open(summary_path, "r", encoding="utf-8") as f:
                    persisted = json.load(f)

        self.assertEqual("relaxed_keyword_v3", snapshot["scoring_version"])
        self.assertEqual(100.0, snapshot["score_percent"])
        self.assertEqual(0, snapshot["failed_count"])
        self.assertEqual([], snapshot["case_items"][0]["hit_detail"]["missing_include"])
        self.assertEqual(["北宋", "大中祥符", "1008-1016"], snapshot["case_items"][0]["hit_detail"]["matched_include"])
        self.assertEqual("relaxed_keyword_v3", persisted["scoring_version"])

    def test_analytics_returns_operation_insights_and_evaluation_summary(self):
        headers = self.admin_headers()
        with patch.object(
            main,
            "get_evaluation_snapshot",
            return_value={
                "ok": True,
                "ready": True,
                "mode": "deepseek",
                "score_percent": 96.8,
                "fact_accuracy": 0.94,
                "failed_count": 3,
                "updated_at": "2026-07-11 10:00:00",
            },
        ):
            response = self.client.get("/api/analytics", headers=headers)

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertIn("operation_insights", data)
        self.assertIn("risk_alerts", data)
        self.assertIn("recommended_actions", data)
        self.assertIn("evaluation_summary", data)
        self.assertTrue(data["recommended_actions"])

    def test_admin_analytics_ai_analysis_calls_deepseek_helper(self):
        headers = self.admin_headers()
        fake_analysis = {
            "provider": "deepseek",
            "summary": "游客集中关注路线与演出，建议把入园后前 30 分钟的导览触点前置。",
            "focus_points": ["路线咨询集中", "演出时间需要前置提醒"],
            "risks": ["低分题会影响导览可信度"],
            "actions": ["首页增加演出时间提示", "补充亲子路线运营话术"],
        }
        with patch.object(main, "generate_operation_ai_analysis", return_value=fake_analysis) as mocked:
            response = self.client.post("/api/admin/analytics/ai-analysis", headers=headers)

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertEqual(fake_analysis, data["analysis"])
        self.assertIn("analytics_snapshot", data)
        mocked.assert_called_once()

    def test_admin_analytics_ai_analysis_accepts_get_for_cached_frontend_compatibility(self):
        headers = self.admin_headers()
        fake_analysis = {
            "provider": "deepseek",
            "summary": "兼容旧前端触发方式，避免运营建议页显示 405。",
            "focus_points": ["游客关注路线"],
            "risks": [],
            "actions": ["继续跟踪热门问题"],
        }
        with patch.object(main, "generate_operation_ai_analysis", return_value=fake_analysis):
            response = self.client.get("/api/admin/analytics/ai-analysis", headers=headers)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(fake_analysis, response.get_json()["analysis"])

    def test_routes_include_contextual_recommendation_reason(self):
        response = self.client.get(
            "/api/routes?interest=亲子家庭&weather=雨&arrival_period=afternoon&companions=老人孩子&stamina=low"
        )

        self.assertEqual(response.status_code, 200)
        routes = response.get_json()["routes"]
        self.assertTrue(routes)
        self.assertIn("recommendation_reason", routes[0])
        self.assertIn("recommendation_context", routes[0])
        reason = routes[0]["recommendation_reason"]
        self.assertTrue(any(word in reason for word in ["亲子", "雨", "室内", "老人", "下午"]))


if __name__ == "__main__":
    unittest.main()
